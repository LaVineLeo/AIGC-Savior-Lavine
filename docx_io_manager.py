# -*- coding: utf-8 -*-
"""
Docx 原文母板重铸与注装脚本 (Docx Template & IO Manager)
提供防 AIGC 的全工作流程，在导出环节能够将改写后的纯文本根据原文字数或段落布局，重新倒模写进旧文档内（保留旧文档内置的大表、格式及其他骨架内容不受波及）。
依赖包: python-docx
"""
import argparse
import os
import sys

try:
    import docx
    from docx.oxml.ns import qn
except ImportError:
    print("❌ 缺失 python-docx 核心驱动，执行前请装载: `pip install python-docx`", file=sys.stderr)
    sys.exit(1)


def extract_docx(doc_path, txt_out_path):
    """
    抽取纯净行文字给 Agent 加工用。
    """
    if not os.path.exists(doc_path):
        print(f"❌ 查无此母版文档: {doc_path}")
        return

    doc = docx.Document(doc_path)
    output_lines = []
    
    # 逐段落扫描
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            # 加入有效段落
            output_lines.append(txt)

    with open(txt_out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(output_lines))
    print(f"✅ 源文档段落抽取完毕（仅抽出了含有实质连绵文本的体量：{len(output_lines)} 组），存放于: {txt_out_path}")


def export_reinject_docx(ref_doc_path, enhanced_txt_path, final_out_path):
    """
    按母板遍历装回法：
    如果从 enhanced_txt 中获得了多行数据，
    按序替换 ref_doc_path 原有的有实质性文字的地方（保证所有未被 LLM 取出碰及的非文字部件与首尾不受伤害）。
    """
    if not os.path.exists(ref_doc_path) or not os.path.exists(enhanced_txt_path):
        print("❌ 核心溯源材料或结晶改写成品缺失，无法对账封装！")
        return

    # 装载刚刚被剥离清洗拉皮后的成品纯文本文段列表
    with open(enhanced_txt_path, "r", encoding="utf-8") as f:
        # 清理由于模型偶发产生的空行断层
        enhanced_lines = [l.strip() for l in f.read().splitlines() if l.strip()]

    doc = docx.Document(ref_doc_path)

    # 中文文字学术字体底层保护
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun') # 尊重用户的特殊默认中文排印

    line_idx = 0
    total_enhanced_len = len(enhanced_lines)

    for para in doc.paragraphs:
        if para.text.strip():
            # 遇到了原始文件中非空的占位段落结构，强行打注新血肉
            if line_idx < total_enhanced_len:
                # 为了格式规整，强抽旧的所有碎裂 run 然后挂上唯一的新的连串 run
                para.clear()
                run = para.add_run(enhanced_lines[line_idx])
                
                # 同步中文基础渲染底架防失效
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                
                line_idx += 1
            else:
                # 新文本比老段落少，把原文这部分留空或者清空，避免文段接续怪乱
                para.clear()

    # 如果有超出的新段落没有排进去？作为新段强制顺手加在底部以绝后患。
    while line_idx < total_enhanced_len:
         new_p = doc.add_paragraph()
         run = new_p.add_run(enhanced_lines[line_idx])
         run.font.name = 'Times New Roman'
         run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
         line_idx += 1

    doc.save(final_out_path)
    print(f"✅ 结晶重铸闭环。基于 '{ref_doc_path}' 原装母本骨架，我们成功覆写注入了脱检测成果，封印为了: {final_out_path}")


def main():
    parser = argparse.ArgumentParser(description="学术版式继承并覆写外设工具 (Reverse Injection Manager)")
    parser.add_argument('--mode', choices=['extract', 'export'], required=True, help="操作形式")
    parser.add_argument('--ref', type=str, required=True, help="原始被提取的母机文档(.docx)")
    parser.add_argument('--text', type=str, help="从大模型手下接过来或者将要发给它的中间态带标点纯文本(.txt)")
    parser.add_argument('--out', type=str, help="成品封装输出档或者待洗中继器")

    args = parser.parse_args()

    if args.mode == "extract":
        extract_docx(args.ref, args.out)
    elif args.mode == "export":
        if not args.text or not args.out:
            print("❌ 执行【全量导出式装配】时必须附带 --text 以及 --out 路径！")
            return
        export_reinject_docx(args.ref, args.text, args.out)

if __name__ == "__main__":
    main()
